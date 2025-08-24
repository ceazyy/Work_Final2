import os
import json
import yaml
import re
from pathlib import Path
from docx import Document
from openai import AzureOpenAI
from typing import Dict, List, Any

class DocumentProcessor:
    def __init__(self, config_path: str = 'config.yaml'):
        """Initialize the document processor with configuration."""
        self.config = self._load_config(config_path)
        self.client = AzureOpenAI(
            api_key=self.config['azure_openai']['api_key'],
            api_version=self.config['azure_openai']['api_version'],
            azure_endpoint=self.config['azure_openai']['endpoint']
        )
        self.input_data = {}
        
    def _load_config(self, config_path: str) -> Dict[str, Any]:
        """Load configuration from YAML file."""
        with open(config_path, 'r', encoding='utf-8') as f:
            config = yaml.safe_load(f)
        
        # Validate Azure OpenAI configuration
        required_keys = ['api_key', 'api_version', 'endpoint', 'deployment_name']
        if 'azure_openai' not in config:
            raise ValueError("azure_openai section missing from config.yaml")
        
        for key in required_keys:
            if key not in config['azure_openai']:
                raise ValueError(f"Missing required key '{key}' in azure_openai config")
            
            if config['azure_openai'][key].startswith('YOUR_'):
                print(f"⚠️  Warning: Please update {key} in config.yaml with your actual Azure OpenAI {key}")
        
        return config
    
    def parse_input_folder(self, input_folder: str = 'input') -> Dict[str, Any]:
        """Parse all files in the input folder and extract information."""
        input_path = Path(input_folder)
        parsed_data = {}
        
        for file_path in input_path.glob('*'):
            if file_path.is_file():
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                    parsed_data[file_path.name] = {
                        'content': content,
                        'path': str(file_path),
                        'size': file_path.stat().st_size,
                        'type': file_path.suffix
                    }
                except Exception as e:
                    parsed_data[file_path.name] = {
                        'error': f'Could not read file: {str(e)}',
                        'path': str(file_path)
                    }
        
        return parsed_data
    
    def extract_document_content(self, template_path: str = 'templates/template.docx') -> Dict[str, Any]:
        """Extract content from the Word document."""
        doc = Document(template_path)
        
        # Extract paragraphs with placeholder detection
        paragraphs = []
        placeholders = set()
        
        # Define instruction words and phrases that indicate placeholders
        instruction_words = [
            'describe', 'list', 'insert', 'define', 'provide', 'explain', 
            'identify', 'specify', 'detail', 'include', 'add', 'enter',
            'fill', 'complete', 'outline', 'summarize', 'document', 'state'
        ]
        
        placeholder_phrases = [
            'purpose and scope', 'architecture specifications', 'block diagram',
            'limitations', 'technical terms', 'abbreviations', 'external interfaces',
            'design decisions', 'assumptions', 'standards', 'protocols',
            'clocking', 'reset', 'power domains'
        ]
        
        for para in doc.paragraphs:
            text = para.text
            paragraphs.append(text)
            
            # Find placeholders in various formats: {{text}}, [text], <text>, etc.
            found_placeholders = re.findall(r'\{\{([^}]+)\}\}|\[([^\]]+)\]|<([^>]+)>|\$\{([^}]+)\}', text)
            
            # Find underscored placeholders like "Title: ____________________"
            underscore_matches = re.findall(r'([^:]*?):\s*_{10,}', text)
            for match in underscore_matches:
                if match.strip():
                    placeholders.add(match.strip())
            
            # Process other placeholder formats
            for groups in found_placeholders:
                for group in groups:
                    if group:
                        placeholders.add(group.strip())
            
            # NEW: Detect instruction-style placeholders
            if text and len(text.strip()) > 10:  # Only process substantial text
                text_lower = text.lower()
                
                # Check if it's an instruction sentence (ends with period, starts with action word)
                is_instruction = (
                    text.strip().endswith('.') and
                    any(text_lower.startswith(word) for word in instruction_words)
                )
                
                # Check if it contains common placeholder phrases
                has_placeholder_phrase = any(phrase in text_lower for phrase in placeholder_phrases)
                
                # Check if it's likely a template instruction
                is_template_instruction = (
                    ('design' in text_lower or 'rtl' in text_lower) and
                    len(text.split()) < 20  # Not too long
                )
                
                if is_instruction or has_placeholder_phrase or is_template_instruction:
                    # Use the full text as the placeholder key
                    placeholders.add(text.strip())
        
        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text
                    row_data.append(cell_text)
                    
                    # Check for placeholders in table cells
                    found_placeholders = re.findall(r'\{\{([^}]+)\}\}|\[([^\]]+)\]|<([^>]+)>|\$\{([^}]+)\}', cell_text)
                    
                    # Check for underscored placeholders in tables
                    underscore_matches = re.findall(r'([^:]*?):\s*_{10,}', cell_text)
                    for match in underscore_matches:
                        if match.strip():
                            placeholders.add(match.strip())
                    
                    for groups in found_placeholders:
                        for group in groups:
                            if group:
                                placeholders.add(group.strip())
                    
                    # NEW: Also check for instruction-style placeholders in table cells
                    if cell_text and len(cell_text.strip()) > 10:
                        text_lower = cell_text.lower()
                        
                        is_instruction = (
                            cell_text.strip().endswith('.') and
                            any(text_lower.startswith(word) for word in instruction_words)
                        )
                        
                        has_placeholder_phrase = any(phrase in text_lower for phrase in placeholder_phrases)
                        
                        is_template_instruction = (
                            ('design' in text_lower or 'rtl' in text_lower) and
                            len(cell_text.split()) < 20
                        )
                        
                        if is_instruction or has_placeholder_phrase or is_template_instruction:
                            placeholders.add(cell_text.strip())
                
                table_data.append(row_data)
            tables.append(table_data)
        
        return {
            'paragraphs': paragraphs,
            'tables': tables,
            'placeholders': list(placeholders),
    'metadata': {
        'author': doc.core_properties.author,
        'created': str(doc.core_properties.created)
    }
}
    
    def fill_placeholder_with_llm(self, placeholder: str, context_data: Dict[str, Any]) -> str:
        """Use LLM to fill a placeholder with relevant content from context data."""
        # Prepare context summary for LLM
        context_summary = "Available information:\n"
        for filename, file_data in context_data.items():
            if 'content' in file_data:
                context_summary += f"\n{filename} ({file_data.get('type', 'unknown')}): {file_data['content'][:500]}..."
            elif 'error' in file_data:
                context_summary += f"\n{filename}: {file_data['error']}"
        
        prompt = f"""You are helping to fill a placeholder in a document template.
        
Placeholder to fill: "{placeholder}"
        
{context_summary}
        
Based on the available information, provide appropriate content to fill this placeholder. 
If you cannot find relevant information, explain why and suggest what information would be needed.
Keep the response concise and directly relevant to the placeholder name.
        """
        
        try:
            response = self.client.chat.completions.create(
                model=self.config['azure_openai']['deployment_name'],
                messages=[
                    {"role": "system", "content": "You are a helpful assistant that fills document placeholders with relevant information from provided context."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=500,
                temperature=0.3
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            return f"ERROR: Could not generate content for '{placeholder}': {str(e)}"
    
    def process_document(self, template_path: str = 'templates/template.docx', input_folder: str = 'input') -> Dict[str, Any]:
        """Main processing function that combines all steps."""
        # Parse input data
        print("Parsing input folder...")
        self.input_data = self.parse_input_folder(input_folder)
        
        # Extract document content
        print("Extracting document content...")
        doc_content = self.extract_document_content(template_path)
        
        # Fill placeholders using LLM
        print("Filling placeholders with LLM...")
        filled_placeholders = {}
        for placeholder in doc_content['placeholders']:
            print(f"  Filling: {placeholder}")
            filled_content = self.fill_placeholder_with_llm(placeholder, self.input_data)
            filled_placeholders[placeholder] = filled_content
        
        return {
            'input_data': self.input_data,
            'document_structure': doc_content,
            'filled_placeholders': filled_placeholders,
            'processing_summary': {
                'placeholders_found': len(doc_content['placeholders']),
                'placeholders_filled': len(filled_placeholders),
                'input_files_processed': len(self.input_data)
            }
        }
    
    def save_results(self, results: Dict[str, Any], output_path: str = 'output_generated/processed_document.json'):
        """Save the processing results to a JSON file."""
        # Create output directory if it doesn't exist
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(results, f, indent=2, ensure_ascii=False)
        
        print(f"Results saved to: {output_path}")
    
    def generate_filled_document(self, processed_json_path: str = 'output_generated/processed_document.json', 
                                template_path: str = 'templates/template.docx',
                                output_path: str = 'output_generated/filled_document.docx'):
        """Generate a filled Word document from processed JSON data."""
        print("Loading processed data...")
        
        # Load the processed JSON data
        with open(processed_json_path, 'r', encoding='utf-8') as f:
            processed_data = json.load(f)
        
        filled_placeholders = processed_data.get('filled_placeholders', {})
        
        # Load the original template
        print("Loading template document...")
        doc = Document(template_path)
        
        # Replace placeholders in paragraphs
        print("Replacing placeholders in document...")
        replacements_made = 0
        
        for para in doc.paragraphs:
            original_text = para.text
            if original_text:
                replaced = False
                
                # First, check for exact full-text matches (instruction-style placeholders)
                for placeholder, filled_content in filled_placeholders.items():
                    if original_text.strip() == placeholder:
                        # Replace entire paragraph with filled content
                        para.clear()
                        para.add_run(filled_content)
                        replacements_made += 1
                        replaced = True
                        break
                
                if not replaced:
                    # Check for underscored placeholders like "Title: ____________________"
                    for placeholder, filled_content in filled_placeholders.items():
                        pattern = f"{re.escape(placeholder)}:\\s*_{{10,}}"
                        
                        if re.search(pattern, original_text):
                            # Clear the paragraph and add new content
                            para.clear()
                            para.add_run(f"{placeholder}: {filled_content}")
                            replacements_made += 1
                            replaced = True
                            break
                
                if not replaced:
                    # Handle other placeholder formats
                    for placeholder, filled_content in filled_placeholders.items():
                        placeholder_patterns = [
                            f"{{{{ {placeholder} }}}}",
                            f"[{placeholder}]",
                            f"<{placeholder}>",
                            f"${{{placeholder}}}"
                        ]
                        
                        for pattern in placeholder_patterns:
                            if pattern in para.text:
                                para.text = para.text.replace(pattern, filled_content)
                                replacements_made += 1
                                replaced = True
                                break
                        
                        if replaced:
                            break
        
        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    original_text = cell.text
                    if original_text:
                        replaced = False
                        
                        # First, check for exact full-text matches (instruction-style placeholders)
                        for placeholder, filled_content in filled_placeholders.items():
                            if original_text.strip() == placeholder:
                                # Replace entire cell content
                                cell.text = filled_content
                                replacements_made += 1
                                replaced = True
                                break
                        
                        if not replaced:
                            # Check for underscored placeholders
                            for placeholder, filled_content in filled_placeholders.items():
                                pattern = f"{re.escape(placeholder)}:\\s*_{{10,}}"
                                if re.search(pattern, original_text):
                                    cell.text = f"{placeholder}: {filled_content}"
                                    replacements_made += 1
                                    replaced = True
                                    break
                        
                        if not replaced:
                            # Handle other placeholder formats
                            for placeholder, filled_content in filled_placeholders.items():
                                placeholder_patterns = [
                                    f"{{{{ {placeholder} }}}}",
                                    f"[{placeholder}]",
                                    f"<{placeholder}>",
                                    f"${{{placeholder}}}"
                                ]
                                
                                for pattern in placeholder_patterns:
                                    if pattern in cell.text:
                                        cell.text = cell.text.replace(pattern, filled_content)
                                        replacements_made += 1
                                        replaced = True
                                        break
                                
                                if replaced:
                                    break
        
        print(f"Total replacements made: {replacements_made}")
        
        # Create output directory if it doesn't exist
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        # Save the filled document
        doc.save(output_path)
        print(f"Filled document saved to: {output_path}")
        
        return {
            'template_path': template_path,
            'processed_json_path': processed_json_path,
            'output_path': output_path,
            'placeholders_replaced': len(filled_placeholders),
            'status': 'success'
        }

def test_azure_openai_config():
    """Test Azure OpenAI configuration without making API calls."""
    try:
        processor = DocumentProcessor()
        config = processor.config['azure_openai']
        
        print("=== Azure OpenAI Configuration ===")
        print(f"API Version: {config['api_version']}")
        print(f"Endpoint: {config['endpoint']}")
        print(f"Deployment Name: {config['deployment_name']}")
        print(f"API Key: {'*' * (len(config['api_key']) - 8) + config['api_key'][-4:] if not config['api_key'].startswith('YOUR_') else 'Not configured'}")
        
        # Check if all values are properly set
        all_configured = all(
            not value.startswith('YOUR_') 
            for value in config.values()
        )
        
        if all_configured:
            print("✅ Configuration appears to be properly set!")
        else:
            print("⚠️  Some configuration values need to be updated")
            
        return all_configured
        
    except Exception as e:
        print(f"❌ Configuration error: {e}")
        return False

def generate_document_from_json(json_path: str = 'output_generated/processed_document.json',
                               template_path: str = 'templates/template.docx',
                               output_path: str = 'output_generated/filled_document.docx'):
    """Standalone function to generate Word document from processed JSON."""
    processor = DocumentProcessor()
    return processor.generate_filled_document(json_path, template_path, output_path)

def main():
    """Main execution function."""
    processor = DocumentProcessor()
    
    # Process the document and extract/fill placeholders
    results = processor.process_document()
    processor.save_results(results)
    
    # Generate the filled Word document
    print("\n" + "="*50)
    print("Generating filled Word document...")
    doc_result = processor.generate_filled_document()
    
    print("\n=== Processing Summary ===")
    summary = results['processing_summary']
    print(f"Input files processed: {summary['input_files_processed']}")
    print(f"Placeholders found: {summary['placeholders_found']}")
    print(f"Placeholders filled: {summary['placeholders_filled']}")
    print(f"Document generated: {doc_result['status']}")
    print(f"Output document: {doc_result['output_path']}")
    
    if results['document_structure']['placeholders']:
        print("\n=== Sample Placeholders and their filled content ===")
        # Show first 5 placeholders as examples
        count = 0
        for placeholder, content in results['filled_placeholders'].items():
            if count >= 5:
                break
            print(f"\n[{placeholder}]:")
            print(f"  {content[:150]}{'...' if len(content) > 150 else ''}")
            count += 1
        
        if len(results['filled_placeholders']) > 5:
            print(f"\n... and {len(results['filled_placeholders']) - 5} more placeholders")

if __name__ == '__main__':
    main()